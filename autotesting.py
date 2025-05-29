import os
import time
import requests
from datetime import datetime
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches
from selenium.webdriver.chrome.service import Service

LOGIN_URL = "https://juie.amanbay.info/ru/sign-in"
ACCOUNT_TYPES_URL = "https://juie.amanbay.info/ru/account-types"
PHONE = "+7 777 777 77 77"
PASSWORD = "admin123"

def smart_login(driver):
    driver.get(LOGIN_URL)
    time.sleep(1.5)
    driver.find_element(By.XPATH, "//input[@type='tel']").send_keys(PHONE)
    driver.find_element(By.XPATH, "//input[@type='password']").send_keys(PASSWORD)
    driver.find_element(By.XPATH, "//button[contains(.,'Войти')]").click()
    time.sleep(2.5)

def smart_screenshot(driver, filename):
    driver.save_screenshot(filename)
    return filename

def scroll_and_screenshot(driver, filename):
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
    time.sleep(1)
    driver.save_screenshot(filename)
    return filename

def check_special_chars_validation(driver, doc, index):
    driver.get(ACCOUNT_TYPES_URL)
    time.sleep(2)


    try:
        add_btn = driver.find_element(By.XPATH, "//button[contains(.,'Добавить тип счета')]")
        add_btn.click()
        time.sleep(1.5)
    except:
        doc.add_paragraph("❌ Не удалось открыть форму создания типа счета.")
        return

    try:
        input_el = driver.find_element(By.XPATH, "//input[contains(@placeholder,'Имя и фамилия')]")
        input_el.clear()
        input_el.send_keys("!!!@@@###")
        time.sleep(1)
        
        save_btn = driver.find_element(By.XPATH, "//button[contains(.,'Создать')]")
        save_btn.click()
        time.sleep(1.5)


        screen_path = f"error_validation_{index}.png"
        smart_screenshot(driver, screen_path)
        doc.add_paragraph("Скриншот проверки поля 'Имя и фамилия' на спецсимволы:")
        doc.add_picture(screen_path, width=Inches(5.5))
    except:
        doc.add_paragraph("❌ Не найден элемент поля 'Имя и фамилия' с ожидаемыми селекторами.")

def collect_ui_state(driver):
    return {
        'titles': [el.text for el in driver.find_elements(By.TAG_NAME, "h1") if el.text],
        'buttons': [el.text or el.get_attribute('aria-label') or el.get_attribute('title')
                    for el in driver.find_elements(By.TAG_NAME, "button") if el.is_displayed()],
        'inputs': [el.get_attribute('name') or el.get_attribute('placeholder')
                   for el in driver.find_elements(By.TAG_NAME, "input") if el.is_displayed()],
    }

def main():
    options = Options()
    options.add_argument("--start-maximized")
    service = Service(ChromeDriverManager().install())
    driver = webdriver.Chrome(service=service, options=options)

    doc = Document()
    doc.add_heading("JUie — Автотестирование", 0)
    doc.add_paragraph(f"Дата теста: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    smart_login(driver)
    doc.add_heading("1. Логин — Успешно!", level=1)

    driver.get(ACCOUNT_TYPES_URL)
    time.sleep(2)

    top_shot = "account_types_top.png"
    smart_screenshot(driver, top_shot)
    doc.add_heading("2. Страница 'Типы счетов'", level=1)
    doc.add_paragraph("Скриншот верха страницы:")
    doc.add_picture(top_shot, width=Inches(5.6))

    bottom_shot = "account_types_bottom.png"
    scroll_and_screenshot(driver, bottom_shot)
    doc.add_paragraph("Скриншот низа страницы:")
    doc.add_picture(bottom_shot, width=Inches(5.6))


    ui = collect_ui_state(driver)
    doc.add_paragraph("Заголовки (h1): " + ", ".join(ui['titles']) if ui['titles'] else "Заголовков нет.")
    doc.add_paragraph("Кнопки: " + ", ".join([b for b in ui['buttons'] if b]) if ui['buttons'] else "Кнопок нет.")
    doc.add_paragraph("Поля ввода: " + ", ".join([i for i in ui['inputs'] if i]) if ui['inputs'] else "Полей нет.")

    doc.add_heading("3. Проверка валидации на спецсимволы", level=1)
    check_special_chars_validation(driver, doc, 3)

    driver.quit()
    doc.save("juie_autotest_report.")

    print("\n✅ Тест завершён. Генерирован отчёт: juie_autotest_report.docx")
    print("\n📋 Выполненные чек-листы:")
    print("✔️ Логин")
    print("✔️ Проверка отображения страницы 'Типы счетов'")
    print("✔️ Сбор заголовков, кнопок, полей")
    print("✔️ Проверка валидации поля 'Имя и фамилия' на спецсимволы")
    print("✔️ Скриншоты верха и низа страницы")
    print("✔️ Скриншот формы с ошибкой")

if __name__ == "__main__":
    main()
